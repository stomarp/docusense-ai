"""
DocuSense AI - Main FastAPI Application

This file:
- Handles file uploads
- Extracts document text
- Performs rule-based compliance analysis
- Runs ML section classification
- Saves documents and reports to PostgreSQL
"""

from fastapi import FastAPI, UploadFile, File, HTTPException, Depends
from pathlib import Path
import uuid

# DOCX reader (aliased to avoid name collision with DB model)
from docx import Document as DocxDocument

# PDF reader
from pypdf import PdfReader

# SQLAlchemy
from sqlalchemy.orm import Session

# Database setup
from backend.app.db.database import Base, engine, get_db
from backend.app.db.models import Document as DBDocument, AnalysisReport

# ML classifier
from backend.app.ml.section_classifier import SectionClassifier
from backend.app.services.risk_intelligence import build_risk_intelligence_report
from backend.app.services.ai_review import build_ai_review


# -------------------------------------------------------
# App Initialization
# -------------------------------------------------------

app = FastAPI(
    title="DocuSense AI",
    description="AI document risk analyzer backend for HR policies, leases, contracts, and compliance documents",
    version="0.1.0"
)

# Create DB tables if not present
Base.metadata.create_all(bind=engine)

# Load ML model once at startup
section_classifier = SectionClassifier()
try:
    section_classifier.load()
except Exception as e:
    print(f"[ML] Section classifier not loaded: {e}")


# -------------------------------------------------------
# File Storage Setup
# -------------------------------------------------------

UPLOAD_DIR = Path("backend/uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


# -------------------------------------------------------
# Text Extraction Helpers
# -------------------------------------------------------

def extract_text_from_docx(file_path: Path) -> str:
    """Extract text from DOCX file."""
    doc = DocxDocument(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def extract_text_from_pdf(file_path: Path) -> str:
    """Extract text from text-based PDF."""
    reader = PdfReader(str(file_path))
    pages_text = []

    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages_text.append(page_text)

    return "\n".join(pages_text)


def split_into_chunks(text: str) -> list[str]:
    """Split document into paragraph chunks for ML."""
    return [line.strip() for line in text.split("\n") if line.strip()]


# -------------------------------------------------------
# Rule-Based Logic
# -------------------------------------------------------

REQUIRED_SECTIONS = {
    "code_of_conduct": {
        "label": "Code of Conduct",
        "keywords": ["code of conduct", "conduct policy"]
    },
    "anti_harassment": {
        "label": "Anti-Harassment Policy",
        "keywords": ["anti-harassment", "harassment", "equal opportunity"]
    },
    "leave_policy": {
        "label": "Leave Policy",
        "keywords": ["leave policy", "vacation", "sick leave", "pto"]
    },
    "working_hours": {
        "label": "Working Hours & Attendance",
        "keywords": ["working hours", "attendance", "timekeeping"]
    },
    "disciplinary_action": {
        "label": "Disciplinary Action Policy",
        "keywords": ["disciplinary", "termination", "corrective action"]
    },
}

RISKY_PHRASES = [
    "may",
    "might",
    "as needed",
    "at management discretion",
    "reasonable",
]


def normalize_text(text: str) -> str:
    return text.lower()


def find_missing_sections(text: str):
    lower_text = normalize_text(text)
    missing = []

    for key, data in REQUIRED_SECTIONS.items():
        if not any(keyword in lower_text for keyword in data["keywords"]):
            missing.append({
                "key": key,
                "label": data["label"],
                "severity": "HIGH"
            })

    return missing


def count_risky_phrases(text: str):
    lower_text = normalize_text(text)
    return {
        phrase: lower_text.count(phrase)
        for phrase in RISKY_PHRASES
        if lower_text.count(phrase) > 0
    }


def calculate_risk_score(missing_sections, risky_phrases):
    score = len(missing_sections) * 20
    score += sum(risky_phrases.values())
    return min(score, 100)


DOCUMENT_TYPE_KEYWORDS = {
    "hr_policy": [
        "employee", "handbook", "code of conduct", "harassment",
        "leave policy", "pto", "attendance", "disciplinary"
    ],
    "lease_agreement": [
        "tenant", "landlord", "rent", "security deposit",
        "premises", "lease", "eviction", "utilities"
    ],
    "offer_letter": [
        "offer", "salary", "compensation", "start date",
        "employment", "benefits", "position", "bonus"
    ],
    "contract": [
        "agreement", "services", "payment terms", "termination",
        "liability", "indemnification", "confidentiality", "governing law"
    ],
    "compliance_document": [
        "compliance", "audit", "regulatory", "procedure",
        "controls", "policy", "monitoring", "reporting"
    ],
}


def detect_document_type(text: str) -> str:
    lower_text = normalize_text(text)
    scores = {}

    for doc_type, keywords in DOCUMENT_TYPE_KEYWORDS.items():
        scores[doc_type] = sum(1 for keyword in keywords if keyword in lower_text)

    best_type = max(scores, key=scores.get)

    if scores[best_type] == 0:
        return "general_document"

    return best_type


def build_plain_english_summary(
    text: str,
    document_type: str,
    risk_score: int,
    missing_sections: list,
    risky_phrases: dict
) -> str:
    risk_level = "low"
    if risk_score >= 70:
        risk_level = "high"
    elif risk_score >= 35:
        risk_level = "medium"

    first_lines = [line.strip() for line in text.split("\n") if line.strip()]
    preview = first_lines[0] if first_lines else "The document was processed successfully."

    return (
        f"This appears to be a {document_type.replace('_', ' ')}. "
        f"The current risk level is {risk_level} with a score of {risk_score}/100. "
        f"DocuSense found {len(missing_sections)} missing section(s) and "
        f"{sum(risky_phrases.values()) if risky_phrases else 0} vague or risky phrase occurrence(s). "
        f"Document preview: {preview[:220]}"
    )


def extract_key_obligations(text: str) -> list[dict]:
    obligation_keywords = [
        "must", "shall", "required", "responsible for", "agree to",
        "obligation", "pay", "provide", "maintain", "comply"
    ]

    lines = [line.strip() for line in text.split("\n") if line.strip()]
    obligations = []

    for line in lines:
        lower_line = normalize_text(line)
        if any(keyword in lower_line for keyword in obligation_keywords):
            obligations.append({
                "text": line[:300],
                "source": "keyword_match"
            })

    return obligations[:8]


def build_red_flags(missing_sections: list, risky_phrases: dict) -> list[dict]:
    red_flags = []

    for section in missing_sections:
        red_flags.append({
            "type": "missing_section",
            "severity": section.get("severity", "HIGH"),
            "message": f"Missing expected section: {section.get('label', section.get('key', 'Unknown section'))}"
        })

    for phrase, count in risky_phrases.items():
        severity = "MEDIUM" if count >= 3 else "LOW"
        red_flags.append({
            "type": "vague_language",
            "severity": severity,
            "message": f"Phrase '{phrase}' appears {count} time(s), which may create ambiguity."
        })

    return red_flags


def build_suggested_improvements(
    document_type: str,
    missing_sections: list,
    risky_phrases: dict
) -> list[str]:
    suggestions = []

    for section in missing_sections:
        label = section.get("label", section.get("key", "missing section"))
        suggestions.append(f"Add a clear {label} section.")

    if risky_phrases:
        suggestions.append(
            "Replace vague words such as may, might, reasonable, or as needed with clearer requirements."
        )

    if document_type == "lease_agreement":
        suggestions.append("Clarify rent, deposit, maintenance, termination, and late fee responsibilities.")
    elif document_type == "offer_letter":
        suggestions.append("Clarify compensation, start date, benefits, employment type, and termination terms.")
    elif document_type == "contract":
        suggestions.append("Clarify payment terms, liability, confidentiality, termination, and dispute resolution.")
    elif document_type == "compliance_document":
        suggestions.append("Clarify ownership, audit process, reporting cadence, and escalation steps.")
    else:
        suggestions.append("Add clear ownership, timelines, responsibilities, and enforcement language.")

    return suggestions[:8]


DOCUMENT_TYPE_LABELS = {
    "hr_policy": "HR Policy",
    "lease_agreement": "Lease Agreement",
    "offer_letter": "Offer Letter",
    "contract": "Contract",
    "compliance_document": "Compliance Document",
    "general_document": "General Document",
}


def get_document_type_label(document_type: str) -> str:
    return DOCUMENT_TYPE_LABELS.get(document_type, "General Document")


def get_risk_level(risk_score: int) -> str:
    if risk_score >= 75:
        return "Critical"
    if risk_score >= 50:
        return "High"
    if risk_score >= 25:
        return "Medium"
    return "Low"


def get_review_priority(risk_score: int, red_flags: list[dict]) -> str:
    high_flags = [flag for flag in red_flags if flag.get("severity") == "HIGH"]

    if risk_score >= 75 or high_flags:
        return "Immediate legal or HR review recommended"
    if risk_score >= 50:
        return "Review before approval"
    if risk_score >= 25:
        return "Needs cleanup before final use"

    return "Low priority review"


def build_scoring_breakdown(
    risk_score: int,
    missing_sections: list[dict],
    risky_phrases: dict
) -> dict:
    missing_section_penalty = len(missing_sections) * 20
    vague_language_penalty = sum(risky_phrases.values()) if risky_phrases else 0
    quality_score = max(0, 100 - risk_score)

    drivers = []

    if missing_sections:
        drivers.append({
            "factor": "Missing expected sections",
            "impact": missing_section_penalty,
            "explanation": "Important sections are missing or not clearly labeled."
        })

    if risky_phrases:
        drivers.append({
            "factor": "Vague or discretionary language",
            "impact": vague_language_penalty,
            "explanation": "Ambiguous wording can create uncertainty or enforcement risk."
        })

    if not drivers:
        drivers.append({
            "factor": "No major rule-based gaps detected",
            "impact": 0,
            "explanation": "The document includes the expected baseline sections for this analyzer version."
        })

    return {
        "risk_score": risk_score,
        "quality_score": quality_score,
        "missing_section_penalty": missing_section_penalty,
        "vague_language_penalty": vague_language_penalty,
        "drivers": drivers
    }


def build_issue_counts(red_flags: list[dict]) -> dict:
    counts = {
        "total": len(red_flags),
        "high": 0,
        "medium": 0,
        "low": 0,
        "by_type": {}
    }

    for flag in red_flags:
        severity = flag.get("severity", "LOW").lower()
        flag_type = flag.get("type", "unknown")

        if severity in counts:
            counts[severity] += 1

        counts["by_type"][flag_type] = counts["by_type"].get(flag_type, 0) + 1

    return counts


def build_top_risks(red_flags: list[dict]) -> list[dict]:
    severity_rank = {
        "HIGH": 3,
        "MEDIUM": 2,
        "LOW": 1
    }

    sorted_flags = sorted(
        red_flags,
        key=lambda flag: severity_rank.get(flag.get("severity", "LOW"), 0),
        reverse=True
    )

    return sorted_flags[:5]


def build_action_plan(
    missing_sections: list[dict],
    risky_phrases: dict,
    suggested_improvements: list[str]
) -> list[dict]:
    actions = []

    for section in missing_sections:
        label = section.get("label", section.get("key", "missing section"))
        actions.append({
            "priority": "High",
            "action": f"Add or rewrite the {label} section.",
            "reason": "Missing sections create avoidable compliance and interpretation risk."
        })

    if risky_phrases:
        actions.append({
            "priority": "Medium",
            "action": "Replace vague wording with specific responsibilities, timelines, and approval rules.",
            "reason": "Ambiguous language makes the document harder to enforce consistently."
        })

    for suggestion in suggested_improvements:
        if not any(existing["action"] == suggestion for existing in actions):
            actions.append({
                "priority": "Medium",
                "action": suggestion,
                "reason": "Improves clarity, completeness, and reviewer confidence."
            })

    return actions[:8]


def build_product_insights(
    text: str,
    document_type: str,
    missing_sections: list[dict],
    risky_phrases: dict
) -> dict:
    words = [word for word in text.split() if word.strip()]
    risk_level = "strong" if not missing_sections and len(risky_phrases) <= 3 else "needs_review"

    return {
        "detected_document_type": document_type,
        "document_type_label": get_document_type_label(document_type),
        "word_count": len(words),
        "character_count": len(text),
        "readiness_status": risk_level,
        "has_missing_sections": bool(missing_sections),
        "has_vague_language": bool(risky_phrases),
        "reviewer_note": (
            "This document is ready for a light review."
            if risk_level == "strong"
            else "This document should be reviewed before it is finalized or signed."
        )
    }


def build_reviewer_verdict(
    risk_score: int,
    missing_sections: list[dict],
    red_flags: list[dict]
) -> str:
    if risk_score >= 75:
        return "High-risk document. Do not approve without expert review."
    if risk_score >= 50:
        return "Material issues detected. Review and revise before use."
    if missing_sections:
        return "Important sections are missing. Add them before final use."
    if red_flags:
        return "Mostly usable, but wording should be tightened before final use."

    return "No major issues detected by the current analyzer."


def build_dashboard_report(
    document_type: str,
    summary: str,
    risk_score: int,
    missing_sections: list[dict],
    key_obligations: list[dict],
    red_flags: list[dict],
    suggested_improvements: list[str],
    risky_phrases: dict,
    text: str
) -> dict:
    risk_level = get_risk_level(risk_score)
    quality_score = max(0, 100 - risk_score)

    return {
        "executive_summary": {
            "document_type": get_document_type_label(document_type),
            "risk_level": risk_level,
            "risk_score": risk_score,
            "quality_score": quality_score,
            "summary": summary,
            "reviewer_verdict": build_reviewer_verdict(
                risk_score,
                missing_sections,
                red_flags
            )
        },
        "risk_overview": {
            "review_priority": get_review_priority(risk_score, red_flags),
            "issue_counts": build_issue_counts(red_flags),
            "scoring_breakdown": build_scoring_breakdown(
                risk_score,
                missing_sections,
                risky_phrases
            )
        },
        "findings": {
            "missing_sections": missing_sections,
            "top_risks": build_top_risks(red_flags),
            "key_obligations": key_obligations,
            "red_flags": red_flags
        },
        "recommendations": {
            "suggested_improvements": suggested_improvements,
            "action_plan": build_action_plan(
                missing_sections,
                risky_phrases,
                suggested_improvements
            )
        },
        "product_insights": build_product_insights(
            text,
            document_type,
            missing_sections,
            risky_phrases
        )
    }


def build_ml_summary(ml_sections: list[dict]) -> list[dict]:
    """
    Group ML predictions by label and compute a small summary:
    - count
    - average confidence
    - top example snippet
    """
    grouped: dict[str, list[dict]] = {}

    for item in ml_sections:
        label = item["label"]
        grouped.setdefault(label, []).append(item)

    summary = []
    for label, items in grouped.items():
        avg_conf = sum(x["confidence"] for x in items) / len(items)
        top_item = max(items, key=lambda x: x["confidence"])

        summary.append({
            "label": label,
            "count": len(items),
            "avg_confidence": round(avg_conf, 4),
            "top_example": top_item["text_preview"],
            "top_confidence": top_item["confidence"],
        })

    # Sort by avg confidence desc
    summary.sort(key=lambda x: x["avg_confidence"], reverse=True)
    return summary


# -------------------------------------------------------
# Health Endpoint
# -------------------------------------------------------

@app.get("/health")
def health_check():
    return {"status": "DocuSense AI is running"}


# -------------------------------------------------------
# Upload Endpoint
# -------------------------------------------------------

@app.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    unique_id = uuid.uuid4().hex
    original_name = file.filename or "uploaded_file"
    ext = Path(original_name).suffix

    stored_filename = f"{unique_id}{ext}"
    stored_path = UPLOAD_DIR / stored_filename

    contents = await file.read()
    stored_path.write_bytes(contents)

    # Save to DB
    doc = DBDocument(
        original_filename=original_name,
        stored_filename=stored_filename,
        file_type=ext.lstrip(".") if ext else "unknown",
    )

    db.add(doc)
    db.commit()
    db.refresh(doc)

    return {
        "message": "File uploaded and saved",
        "document_id": doc.id,
        "stored_filename": stored_filename
    }


@app.post("/extract-text/{stored_filename}")
def extract_text(stored_filename: str):
    """
    Extract text from a previously uploaded file.

    For now, we support:
    - .docx (Word documents)
    - .txt  (plain text)

    Later we will add:
    - PDF extraction
    """
    file_path = UPLOAD_DIR / stored_filename

    # If file does not exist, return a 404 error
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    # Detect file type by extension
    ext = file_path.suffix.lower()

    if ext == ".docx":
      text = extract_text_from_docx(file_path)
    elif ext == ".pdf":
      text = extract_text_from_pdf(file_path)
    elif ext == ".txt":
      text = file_path.read_text(encoding="utf-8", errors="ignore")
    else:
      raise HTTPException(
        status_code=400,
        detail="Unsupported file type for text extraction (supported: .docx, .pdf, .txt)"
    )

    # Return extracted text (we'll store it in DB in next steps)
    return {
        "stored_filename": stored_filename,
        "characters": len(text),
        "preview": text[:500],   # show only first 500 chars
        "text": text             # full text (ok for now; later we may not return full)
    }


# -------------------------------------------------------
# Analyze Endpoint
# -------------------------------------------------------

@app.post("/analyze/{stored_filename}")
def analyze_document(
    stored_filename: str,
    db: Session = Depends(get_db)
):
    file_path = UPLOAD_DIR / stored_filename

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    ext = file_path.suffix.lower()

    if ext == ".docx":
        text = extract_text_from_docx(file_path)
    elif ext == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif ext == ".txt":
        text = file_path.read_text(encoding="utf-8", errors="ignore")
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    # Rule-based analysis
    missing_sections = find_missing_sections(text)
    risky_phrases = count_risky_phrases(text)
    risk_score = calculate_risk_score(missing_sections, risky_phrases)

    # DocuSense report fields
    document_type = detect_document_type(text)
    summary = build_plain_english_summary(
        text=text,
        document_type=document_type,
        risk_score=risk_score,
        missing_sections=missing_sections,
        risky_phrases=risky_phrases,
    )
    key_obligations = extract_key_obligations(text)
    red_flags = build_red_flags(missing_sections, risky_phrases)
    suggested_improvements = build_suggested_improvements(
        document_type=document_type,
        missing_sections=missing_sections,
        risky_phrases=risky_phrases,
    )

    risk_level = get_risk_level(risk_score)
    quality_score = max(0, 100 - risk_score)
    review_priority = get_review_priority(risk_score, red_flags)
    scoring_breakdown = build_scoring_breakdown(
        risk_score=risk_score,
        missing_sections=missing_sections,
        risky_phrases=risky_phrases,
    )
    issue_counts = build_issue_counts(red_flags)
    top_risks = build_top_risks(red_flags)
    action_plan = build_action_plan(
        missing_sections=missing_sections,
        risky_phrases=risky_phrases,
        suggested_improvements=suggested_improvements,
    )
    product_insights = build_product_insights(
        text=text,
        document_type=document_type,
        missing_sections=missing_sections,
        risky_phrases=risky_phrases,
    )
    reviewer_verdict = build_reviewer_verdict(
        risk_score=risk_score,
        missing_sections=missing_sections,
        red_flags=red_flags,
    )
    product_report = build_dashboard_report(
        document_type=document_type,
        summary=summary,
        risk_score=risk_score,
        missing_sections=missing_sections,
        key_obligations=key_obligations,
        red_flags=red_flags,
        suggested_improvements=suggested_improvements,
        risky_phrases=risky_phrases,
        text=text,
    )

    # ML Section Classification

    # ML Section Classification (optional)
    ml_sections = []
    ml_summary = []

    CONF_THRESHOLD = 0.35

    if section_classifier.is_ready():
        chunks = split_into_chunks(text)
        raw_predictions = section_classifier.predict_many(chunks)
        ml_sections = [p for p in raw_predictions if p["confidence"] >= CONF_THRESHOLD]
        ml_sections = sorted(
            ml_sections,
            key=lambda x: x["confidence"],
            reverse=True
        )[:10]
        ml_summary = build_ml_summary(ml_sections)
   


    intelligence_report = build_risk_intelligence_report(
        text=text,
        source_filename=stored_filename,
        ml_summary=ml_summary,
        ml_sections=ml_sections,
    )
    ai_review = build_ai_review(intelligence_report)

    # Find document in DB
    doc = db.query(DBDocument).filter(
        DBDocument.stored_filename == stored_filename
    ).first()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found in DB")

    # Save report
    report = AnalysisReport(
        document_id=doc.id,
        risk_score=risk_score,
        missing_sections=missing_sections,
        risky_phrases=risky_phrases,
    )

    db.add(report)
    db.commit()
    db.refresh(report)

    return {
        "document_id": doc.id,
        "report_id": report.id,
        "intelligence_report": intelligence_report,
        "executive_summary": intelligence_report["executive_summary"],
        "document_classification": intelligence_report["document_classification"],
        "risk_scores": intelligence_report["risk_scores"],
        "clause_coverage": intelligence_report["clause_coverage"],
        "risk_findings": intelligence_report["risk_findings"],
        "ai_review": ai_review,
        "ai_review_notes": intelligence_report["ai_review_notes"],
        "document_type": document_type,
        "document_type_label": get_document_type_label(document_type),
        "summary": summary,
        "risk_level": risk_level,
        "risk_score": risk_score,
        "quality_score": quality_score,
        "review_priority": review_priority,
        "reviewer_verdict": reviewer_verdict,
        "scoring_breakdown": scoring_breakdown,
        "issue_counts": issue_counts,
        "top_risks": top_risks,
        "missing_sections": missing_sections,
        "key_obligations": key_obligations,
        "red_flags": red_flags,
        "suggested_improvements": suggested_improvements,
        "action_plan": action_plan,
        "product_insights": product_insights,
        "product_report": product_report,
        "legacy_report": {
            "risk_score": risk_score,
            "missing_sections": missing_sections,
            "risky_phrases": risky_phrases,
        },
        "ml_summary": ml_summary,
        "ml_sections": ml_sections,
        "metadata": {
            "extracted_characters": len(text),
            "analysis_type": "docusense_product_report_v3 + ml_classifier",
            "ml_conf_threshold": CONF_THRESHOLD,
            "ml_sections_returned": len(ml_sections),
        }
    }

    
# -------------------------------------------------------
# Documents + Reports History Endpoints
# -------------------------------------------------------

@app.get("/documents")
def list_documents(db: Session = Depends(get_db)):
    """
    List all uploaded documents.
    For each document, include the most recent report (if it exists).
    """
    docs = db.query(DBDocument).order_by(DBDocument.id.desc()).all()
    results = []

    for doc in docs:
        latest_report = (
            db.query(AnalysisReport)
            .filter(AnalysisReport.document_id == doc.id)
            .order_by(AnalysisReport.id.desc())
            .first()
        )

        results.append({
            "document_id": doc.id,
            "original_filename": doc.original_filename,
            "stored_filename": doc.stored_filename,
            "file_type": doc.file_type,
            "created_at": str(doc.created_at),
            "latest_report": None if not latest_report else {
                "report_id": latest_report.id,
                "risk_score": latest_report.risk_score,
                "created_at": str(latest_report.created_at),
            }
        })

    return results


@app.get("/documents/{document_id}/reports")
def list_reports(document_id: int, db: Session = Depends(get_db)):
    """
    List all analysis reports for a specific document (report history).
    """
    doc = db.query(DBDocument).filter(DBDocument.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    reports = (
        db.query(AnalysisReport)
        .filter(AnalysisReport.document_id == document_id)
        .order_by(AnalysisReport.id.desc())
        .all()
    )

    return {
        "document_id": doc.id,
        "original_filename": doc.original_filename,
        "stored_filename": doc.stored_filename,
        "reports": [
            {
                "report_id": r.id,
                "risk_score": r.risk_score,
                "missing_sections": r.missing_sections,
                "risky_phrases": r.risky_phrases,
                "created_at": str(r.created_at),
            }
            for r in reports
        ],
    }